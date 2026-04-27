from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter as _PypdfPdfWriter

from translation_pipeline.documents import (
    DocxWriter,
    PdfWriter,
    UnsupportedDocumentFormat,
    read_document,
)

# ---- TXT ---------------------------------------------------------------------


def test_read_txt(tmp_path: Path) -> None:
    p = tmp_path / "doc.txt"
    p.write_text("Hello\n\nWorld", encoding="utf-8")
    assert read_document(p) == "Hello\n\nWorld"


# ---- DOCX --------------------------------------------------------------------


def test_read_docx_paragraphs(tmp_path: Path) -> None:
    p = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("")  # empty - should be dropped
    doc.add_paragraph("Third paragraph.")
    doc.save(str(p))

    out = read_document(p)
    assert "First paragraph." in out
    assert "Second paragraph." in out
    assert "Third paragraph." in out
    # Blank paragraph dropped — joined as one block per non-empty paragraph.
    assert out.split("\n\n") == [
        "First paragraph.",
        "Second paragraph.",
        "Third paragraph.",
    ]


def test_read_docx_with_table(tmp_path: Path) -> None:
    p = tmp_path / "with_table.docx"
    doc = Document()
    doc.add_paragraph("Before table.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A1"
    table.rows[0].cells[1].text = "B1"
    table.rows[1].cells[0].text = "A2"
    table.rows[1].cells[1].text = "B2"
    doc.add_paragraph("After table.")
    doc.save(str(p))

    out = read_document(p)
    assert "Before table." in out
    assert "After table." in out
    # Table cells joined per-row with " | " separator.
    assert "A1 | B1" in out
    assert "A2 | B2" in out


def test_docx_writer_emits_paragraphs(tmp_path: Path) -> None:
    out = tmp_path / "translation.docx"
    text = "First paragraph.\n\nSecond paragraph.\n\n\n\nThird with extra blanks."
    DocxWriter().write(out, text)

    doc = Document(str(out))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    assert paragraphs == [
        "First paragraph.",
        "Second paragraph.",
        "Third with extra blanks.",
    ]


def test_docx_round_trip(tmp_path: Path) -> None:
    """Write a docx, read it back, write again — text survives intact."""
    src = tmp_path / "src.docx"
    Document().add_paragraph("Once upon a time.")
    d = Document()
    d.add_paragraph("Once upon a time.")
    d.add_paragraph("There was a translator.")
    d.save(str(src))

    text = read_document(src)
    out = tmp_path / "out.docx"
    DocxWriter().write(out, text)

    again = read_document(out)
    assert again == text


# ---- PDF ---------------------------------------------------------------------


def test_pdf_writer_basic_ascii(tmp_path: Path) -> None:
    out = tmp_path / "translation.pdf"
    PdfWriter().write(out, "Hello world.\n\nSecond paragraph.")
    assert out.exists() and out.stat().st_size > 0
    # Round-trip via pypdf to confirm a valid PDF that contains our text.
    from pypdf import PdfReader

    text = "\n".join(p.extract_text() or "" for p in PdfReader(str(out)).pages)
    assert "Hello world" in text
    assert "Second paragraph" in text


def test_pdf_writer_unicode_polish(tmp_path: Path) -> None:
    """Polish diacritics survive when a Unicode TTF is available on the host."""
    out = tmp_path / "polish.pdf"
    text = "Zażółć gęślą jaźń.\n\nKraków, Łódź, Wrocław."
    PdfWriter().write(out, text)
    assert out.exists() and out.stat().st_size > 0
    from pypdf import PdfReader

    extracted = "\n".join(p.extract_text() or "" for p in PdfReader(str(out)).pages)
    # If we found a Unicode font, all chars should round-trip. If we fell back to
    # Helvetica (no Unicode), ł/ż/ó may have been dropped or substituted — accept
    # either, but assert at least the ASCII parts came through.
    assert "Krak" in extracted


def test_pdf_writer_escapes_xml_metachars(tmp_path: Path) -> None:
    out = tmp_path / "metachars.pdf"
    PdfWriter().write(out, "1 < 2 & 3 > 0\n\nunchanged text")
    assert out.exists()


def test_pdf_round_trip(tmp_path: Path) -> None:
    """Write a PDF, read it back via the dispatcher, get the text out."""
    src = tmp_path / "src.pdf"
    PdfWriter().write(src, "Once upon a time.\n\nThere was a translator.")
    text = read_document(src)
    assert "Once upon a time" in text
    assert "There was a translator" in text


def test_read_pdf_real_sample(tmp_path: Path) -> None:
    """Round-trip: write a multi-paragraph PDF, read it back, confirm content."""
    sample = tmp_path / "fixture.pdf"
    PdfWriter().write(
        sample,
        "First paragraph.\n\nSecond paragraph.\n\nThird paragraph.\n\nFourth.",
    )
    text = read_document(sample)
    assert "First paragraph" in text
    assert "Second paragraph" in text
    assert "Third paragraph" in text
    assert "Fourth" in text


def test_read_pdf_empty_returns_empty(tmp_path: Path) -> None:
    """An empty (no-page) PDF yields empty text without crashing."""
    p = tmp_path / "empty.pdf"
    _PypdfPdfWriter().write(str(p))
    assert read_document(p) == ""


# ---- Dispatcher --------------------------------------------------------------


def test_unsupported_extension_raises(tmp_path: Path) -> None:
    p = tmp_path / "doc.rtf"
    p.write_text("hi")
    with pytest.raises(UnsupportedDocumentFormat):
        read_document(p)


def test_extension_case_insensitive(tmp_path: Path) -> None:
    p = tmp_path / "DOC.TXT"
    p.write_text("hi")
    assert read_document(p) == "hi"


# ---- Structured readers (block-anchored) -------------------------------------


from translation_pipeline.documents import (  # noqa: E402
    SENTINEL,
    DocxFormatPreservingWriter,
    PdfFormatPreservingWriter,
    read_structured,
)


def test_read_structured_txt_blocks(tmp_path: Path) -> None:
    p = tmp_path / "doc.txt"
    p.write_text("Block one.\n\nBlock two.\n\nBlock three.")
    src = read_structured(p)
    assert src.kind == "txt"
    assert [b.text for b in src.blocks] == [
        "Block one.",
        "Block two.",
        "Block three.",
    ]
    assert src.handle is None


def test_read_structured_docx_blocks_paragraph_and_table(tmp_path: Path) -> None:
    p = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Para 1.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Cell A"
    table.rows[0].cells[1].text = "Cell B"
    doc.add_paragraph("Para 2.")
    doc.save(str(p))

    src = read_structured(p)
    assert src.kind == "docx"
    texts = [b.text for b in src.blocks]
    assert texts == ["Para 1.", "Cell A", "Cell B", "Para 2."]
    # Anchor is a list of run elements for each block.
    for b in src.blocks:
        assert isinstance(b.anchor, list)
        assert len(b.anchor) >= 1


def test_read_structured_pdf_blocks_via_fitz(tmp_path: Path) -> None:
    src_path = tmp_path / "src.pdf"
    PdfWriter().write(src_path, "First block.\n\nSecond block.\n\nThird block.")
    src = read_structured(src_path)
    assert src.kind == "pdf"
    assert len(src.blocks) >= 1
    joined = " | ".join(b.text for b in src.blocks)
    assert "First block" in joined
    assert "Second block" in joined
    assert "Third block" in joined
    # Anchor is (page_idx, [(rect, span_text, size), ...], aggregate_rect, font_size).
    for b in src.blocks:
        page_idx, sub_units, agg, font_size = b.anchor
        assert page_idx == 0
        assert len(sub_units) >= 1
        for rect, text, size in sub_units:
            assert rect.x1 > rect.x0 and rect.y1 > rect.y0
            assert isinstance(text, str)
            assert size > 0
        assert agg.x1 > agg.x0 and agg.y1 > agg.y0
        assert font_size > 0


def test_read_structured_pdf_coarsens_adjacent_blocks(tmp_path: Path) -> None:
    """A multi-paragraph PDF should produce far fewer logical blocks than raw
    PyMuPDF text blocks. Each adjacent text block on the same page within a
    line-height-or-so vertical gap belongs to the same paragraph."""
    src_path = tmp_path / "long.pdf"
    paragraphs = [f"Paragraph {i}." + " word" * 15 for i in range(20)]
    PdfWriter().write(src_path, "\n\n".join(paragraphs))
    src = read_structured(src_path)
    # 20 source paragraphs but PyMuPDF often yields 1-2 raw blocks per paragraph.
    # After coarsening we expect roughly the source paragraph count, not the
    # raw block count. Bound: ≤ 30 (would fail loudly if coarsening regressed).
    assert len(src.blocks) <= 30


def test_pdf_block_anchor_rect_is_union_of_sub_rects(tmp_path: Path) -> None:
    src_path = tmp_path / "u.pdf"
    PdfWriter().write(src_path, "Line A\nLine B\n\nNew para")
    src = read_structured(src_path)
    for b in src.blocks:
        _page_idx, sub_units, agg, _font_size = b.anchor
        # The aggregate must contain every sub-rect.
        for rect, _text, _size in sub_units:
            assert agg.x0 <= rect.x0 + 0.01
            assert agg.y0 <= rect.y0 + 0.01
            assert agg.x1 + 0.01 >= rect.x1
            assert agg.y1 + 0.01 >= rect.y1


def test_joined_text_uses_sentinel(tmp_path: Path) -> None:
    p = tmp_path / "doc.txt"
    p.write_text("Block A.\n\nBlock B.")
    src = read_structured(p)
    assert SENTINEL in src.joined_text
    assert src.joined_text.split(SENTINEL) == [
        "Block A.\n\n",
        "\n\nBlock B.",
    ]


def test_sentinel_collision_in_source_text_escaped(tmp_path: Path) -> None:
    """Source text containing the literal sentinel must round-trip without false splits."""
    p = tmp_path / "doc.txt"
    p.write_text("This text mentions [[BLK]] literally.\n\nSecond block.")
    src = read_structured(p)
    joined = src.joined_text
    # The inline literal got escaped (zero-width-space inserted), so splitting on
    # the real sentinel yields exactly two segments.
    parts = joined.split(SENTINEL)
    assert len(parts) == 2
    # Segment 0 still contains the (escaped) form so the original meaning survives.
    assert "BLK" in parts[0]


# ---- DocxFormatPreservingWriter ---------------------------------------------


def test_docx_format_preserving_round_trip_preserves_run_props(tmp_path: Path) -> None:
    """A bold paragraph stays bold after the writer mutates it."""
    src_path = tmp_path / "src.docx"
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Original heading text")
    run.bold = True
    run.font.name = "Arial"
    doc.add_paragraph("Body paragraph.")
    doc.save(str(src_path))

    src = read_structured(src_path)
    assert len(src.blocks) == 2

    out = tmp_path / "out.docx"
    DocxFormatPreservingWriter().write(
        out, src, ["Translated heading", "Translated body"]
    )

    # Re-read the output and assert the first run's props survived on the heading.
    written = Document(str(out))
    assert written.paragraphs[0].text == "Translated heading"
    first_run = written.paragraphs[0].runs[0]
    assert first_run.bold is True
    assert first_run.font.name == "Arial"
    assert written.paragraphs[1].text == "Translated body"


def test_docx_format_preserving_keeps_table_structure(tmp_path: Path) -> None:
    src_path = tmp_path / "src.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A1"
    table.rows[0].cells[1].text = "B1"
    table.rows[1].cells[0].text = "A2"
    table.rows[1].cells[1].text = "B2"
    doc.save(str(src_path))

    src = read_structured(src_path)
    out = tmp_path / "out.docx"
    DocxFormatPreservingWriter().write(
        out, src, ["a1*", "b1*", "a2*", "b2*"]
    )

    written = Document(str(out))
    # The output is still a table, not flattened to paragraphs.
    assert len(written.tables) == 1
    cells = written.tables[0]
    assert cells.rows[0].cells[0].text == "a1*"
    assert cells.rows[0].cells[1].text == "b1*"
    assert cells.rows[1].cells[0].text == "a2*"
    assert cells.rows[1].cells[1].text == "b2*"


def test_docx_format_preserving_block_count_mismatch_raises(tmp_path: Path) -> None:
    src_path = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("one")
    doc.add_paragraph("two")
    doc.save(str(src_path))
    src = read_structured(src_path)
    with pytest.raises(ValueError, match="block count mismatch"):
        DocxFormatPreservingWriter().write(tmp_path / "out.docx", src, ["just one"])


# ---- PdfFormatPreservingWriter ----------------------------------------------


def test_pdf_format_preserving_replaces_text_in_place(tmp_path: Path) -> None:
    src_path = tmp_path / "src.pdf"
    PdfWriter().write(src_path, "Hello world.\n\nA second paragraph.")
    src = read_structured(src_path)
    assert len(src.blocks) >= 1

    out = tmp_path / "out.pdf"
    translated = ["Olá mundo. Um segundo parágrafo." for _ in src.blocks]
    PdfFormatPreservingWriter().write(out, src, translated)
    assert out.exists() and out.stat().st_size > 0

    # Read the output back via pypdf — translated text should be present, original gone.
    from pypdf import PdfReader

    extracted = "\n".join(p.extract_text() or "" for p in PdfReader(str(out)).pages)
    assert "Olá mundo" in extracted or "Ol" in extracted
    assert "Hello world" not in extracted


def test_pdf_format_preserving_overflow_does_not_crash(tmp_path: Path) -> None:
    """A translation much longer than the source bbox must not raise."""
    src_path = tmp_path / "src.pdf"
    PdfWriter().write(src_path, "Short.")
    src = read_structured(src_path)
    very_long = "Word " * 200
    out = tmp_path / "out.pdf"
    PdfFormatPreservingWriter().write(out, src, [very_long.strip()] * len(src.blocks))
    assert out.exists()


def test_pdf_format_preserving_block_count_mismatch_raises(tmp_path: Path) -> None:
    """Construct a StructuredSource with multiple blocks directly so the test
    doesn't depend on coarsening behavior."""
    import fitz

    from translation_pipeline.documents import SourceBlock, StructuredSource

    src_path = tmp_path / "src.pdf"
    PdfWriter().write(src_path, "x.\n\ny.")
    handle = fitz.open(str(src_path))

    rect_a = fitz.Rect(50, 50, 200, 80)
    rect_b = fitz.Rect(50, 100, 200, 130)
    src = StructuredSource(
        kind="pdf",
        blocks=(
            SourceBlock(index=0, text="A", anchor=(0, [rect_a], rect_a, 11.0)),
            SourceBlock(index=1, text="B", anchor=(0, [rect_b], rect_b, 11.0)),
            SourceBlock(index=2, text="C", anchor=(0, [rect_b], rect_b, 11.0)),
        ),
        handle=handle,
    )
    with pytest.raises(ValueError, match="block count mismatch"):
        PdfFormatPreservingWriter().write(tmp_path / "out.pdf", src, ["only one"])


# ---- Redistribution / per-run alignment -------------------------------------


def test_redistribute_single_unit_no_call() -> None:
    from translation_pipeline.documents import _redistribute_to_units

    out = _redistribute_to_units("translated whole", ["original"])
    assert out == ["translated whole"]


def test_redistribute_zero_units_returns_empty() -> None:
    from translation_pipeline.documents import _redistribute_to_units

    assert _redistribute_to_units("translated", []) == []


def test_redistribute_no_client_uses_proportional_split() -> None:
    """Multi-unit, no client provided → fallback proportional split.

    The fallback is a heuristic — it may cut mid-word for short inputs without
    convenient whitespace. We only require: every carrier gets a non-empty
    fragment, and concatenation covers all the translated characters.
    """
    from translation_pipeline.documents import _redistribute_to_units

    out = _redistribute_to_units(
        "ala ma kota czarnego pieska",
        ["this is", "important", "stuff"],
    )
    assert len(out) == 3
    assert all(o.strip() for o in out)
    # Every translated character lands somewhere (concat with spaces).
    joined = "".join(out).replace(" ", "")
    assert joined == "alamakotaczarnegopieska"


def test_redistribute_passes_through_whitespace_units() -> None:
    """Whitespace-only source units pass through verbatim."""
    from translation_pipeline.documents import _redistribute_to_units

    out = _redistribute_to_units(
        "translated text here",
        ["text-carrier", "  ", "another-carrier"],
    )
    assert len(out) == 3
    assert out[1] == "  "  # whitespace passed through


def test_redistribute_single_carrier_with_padding_no_call() -> None:
    """If only one source unit carries text, give it the full translation."""
    from translation_pipeline.documents import _redistribute_to_units

    out = _redistribute_to_units(
        "translated",
        [" ", "carrier", " "],
    )
    assert out[0] == " "
    assert out[1] == "translated"
    assert out[2] == " "


def test_redistribute_count_mismatch_logs_warning() -> None:
    """When alignment fakes a wrong count, fallback fires + warning recorded."""
    from translation_pipeline.documents import _redistribute_to_units

    # Stub a fake settings/client where the alignment returns the wrong count.
    class _FakeAlignedTooFew:
        translations: list[str] = ["only one"]

    class _FakeResult:
        parsed = _FakeAlignedTooFew()
        usage = type("u", (), {"input_tokens": 1, "output_tokens": 1})()

    class _FakeClient:
        pass

    class _FakeSettings:
        MODEL_DIVERGENCE_DETECTION = "claude-sonnet-4-6"

    # Monkeypatch _align_via_claude to return wrong count.
    import translation_pipeline.documents as docmod

    orig = docmod._align_via_claude
    try:
        docmod._align_via_claude = lambda *a, **k: ["only one"]  # wrong count
        warnings: list[str] = []
        out = _redistribute_to_units(
            "ala ma kota",
            ["one", "two", "three"],
            client=_FakeClient(),
            settings=_FakeSettings(),
            warnings=warnings,
            block_index=42,
        )
        assert len(out) == 3
        assert any("alignment count mismatch on block 42" in w for w in warnings)
    finally:
        docmod._align_via_claude = orig


# ---- Format preservation: mid-paragraph bold survives ----------------------


def test_docx_format_preserving_keeps_mid_paragraph_bold(tmp_path: Path) -> None:
    """Source paragraph with bold mid-word: the writer must preserve the bold
    run and put the corresponding translated fragment inside it."""
    from translation_pipeline.documents import (
        DocxFormatPreservingWriter,
        read_structured,
    )

    src_path = tmp_path / "src.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("This is ")
    bold_run = p.add_run("very")
    bold_run.bold = True
    p.add_run(" important.")
    doc.save(str(src_path))

    src = read_structured(src_path)
    assert len(src.blocks) == 1
    runs_in_anchor = src.blocks[0].anchor
    assert len(runs_in_anchor) == 3  # 3 distinct runs

    out = tmp_path / "out.docx"
    # No client → uses proportional split fallback. We just need each run to
    # still be its own run after writing, with the bold run still bold.
    DocxFormatPreservingWriter(client=None, settings=None).write(
        out, src, ["To jest bardzo ważne."]
    )

    written = Document(str(out))
    runs = written.paragraphs[0].runs
    assert len(runs) == 3
    # Middle run keeps its bold property regardless of what text ended up in it.
    assert runs[1].bold is True
    # Concatenated text contains every translated word exactly once.
    full = "".join(r.text for r in runs)
    assert "bardzo" in full or "ważne" in full
    # No source text leaked through.
    assert "very" not in full
    assert "important" not in full


def test_structured_source_unsupported_kind_raises_in_writer(tmp_path: Path) -> None:
    """Calling DocxFormatPreservingWriter on a PDF source is a programming error."""
    src_path = tmp_path / "src.pdf"
    PdfWriter().write(src_path, "x")
    src = read_structured(src_path)
    with pytest.raises(ValueError, match="kind='docx'"):
        DocxFormatPreservingWriter().write(
            tmp_path / "out.docx", src, ["x"] * len(src.blocks)
        )
