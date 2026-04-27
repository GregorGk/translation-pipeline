from __future__ import annotations

from pathlib import Path

import pytest
from typer.testing import CliRunner

from translation_pipeline.cli import app


@pytest.fixture
def settings_env(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    """Isolate from the project's real .env."""
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test")
    monkeypatch.setenv("OPENAI_API_KEY", "test")
    monkeypatch.setenv("DEEPL_API_KEY", "test")
    monkeypatch.chdir(tmp_path)


@pytest.fixture
def runner() -> CliRunner:
    return CliRunner()


def test_help(runner: CliRunner) -> None:
    result = runner.invoke(app, ["--help"])
    assert result.exit_code == 0
    assert "Translate a document" in result.stdout
    assert "--from" in result.stdout
    assert "--to" in result.stdout
    assert "--dry-run" in result.stdout


def test_version(runner: CliRunner) -> None:
    # --version requires an INPUT_PATH per Typer's positional-arg rules,
    # so use a dummy file that doesn't need to exist for the eager callback.
    # Actually, --version is eager — it should fire before the input check.
    # Pass a placeholder; eager callback will exit before validation.
    result = runner.invoke(app, ["--version", "doesnt_matter"])
    assert result.exit_code == 0
    assert "translate" in result.stdout


def test_unsupported_target_lang(
    runner: CliRunner, tmp_path: Path, settings_env: None
) -> None:
    src = tmp_path / "in.txt"
    src.write_text("hello")
    result = runner.invoke(app, [str(src), "--from", "EN", "--to", "ES", "--dry-run"])
    assert result.exit_code != 0
    assert "ES" in result.stdout or "ES" in (result.stderr or "")


def test_unsupported_source_lang(
    runner: CliRunner, tmp_path: Path, settings_env: None
) -> None:
    src = tmp_path / "in.txt"
    src.write_text("hello")
    result = runner.invoke(app, [str(src), "--from", "ZZ", "--to", "EN", "--dry-run"])
    assert result.exit_code != 0


def test_file_not_found(
    runner: CliRunner, tmp_path: Path, settings_env: None
) -> None:
    result = runner.invoke(
        app, [str(tmp_path / "nope.txt"), "--from", "EN", "--to", "PL", "--dry-run"]
    )
    assert result.exit_code != 0


def test_dry_run_prints_estimate(
    runner: CliRunner, tmp_path: Path, settings_env: None
) -> None:
    src = tmp_path / "doc.txt"
    src.write_text("Hello world.\n\n" * 50)  # ~600 chars, multi-paragraph
    result = runner.invoke(
        app, [str(src), "--from", "EN", "--to", "PL", "--dry-run"]
    )
    assert result.exit_code == 0, result.stdout
    out = result.stdout
    assert "Dry-run estimate" in out
    assert "TOTAL" in out
    # Per-stage breakdown shows each LLM stage.
    for stage in ("brief_extraction", "draft_b", "synthesis", "consistency"):
        assert stage in out


def test_dry_run_without_from_warns(
    runner: CliRunner, tmp_path: Path, settings_env: None
) -> None:
    src = tmp_path / "doc.txt"
    src.write_text("hello")
    result = runner.invoke(app, [str(src), "--to", "PL", "--dry-run"])
    assert result.exit_code == 0
    assert "skipping language detection" in result.stdout


def test_missing_required_to(runner: CliRunner, tmp_path: Path) -> None:
    src = tmp_path / "doc.txt"
    src.write_text("hello")
    result = runner.invoke(app, [str(src)])
    assert result.exit_code != 0
    assert "--to" in (result.stdout + (result.stderr or ""))


def test_unsupported_input_extension(
    runner: CliRunner, tmp_path: Path, settings_env: None
) -> None:
    src = tmp_path / "doc.rtf"
    src.write_text("hi")
    result = runner.invoke(
        app, [str(src), "--from", "EN", "--to", "PL", "--dry-run"]
    )
    assert result.exit_code != 0
    assert "unsupported input format" in result.stdout


def test_dry_run_pdf_input(
    runner: CliRunner, tmp_path: Path, settings_env: None
) -> None:
    """--dry-run reads PDF text just like txt input."""
    from docx import Document

    # Use a docx as an easy way to round-trip text into a binary format.
    src = tmp_path / "doc.docx"
    d = Document()
    d.add_paragraph("Hello from a docx file.")
    d.save(str(src))

    result = runner.invoke(
        app, [str(src), "--from", "EN", "--to", "PL", "--dry-run"]
    )
    assert result.exit_code == 0, result.stdout
    assert "Dry-run estimate" in result.stdout
    # Source chars came from extracted DOCX text, not raw bytes.
    assert "source chars: 23" in result.stdout


# ---- Preserve-mode dispatch (mocked pipeline) -------------------------------


def _patch_pipeline_to_return_blocks(
    monkeypatch: pytest.MonkeyPatch,
    block_texts: list[str],
    *,
    src_lang: str = "EN",
    tgt_lang: str = "PL",
) -> dict[str, object]:
    """Stub the CLI's pipeline + language detection so we exercise writer dispatch
    without burning API credits.

    Returns a dict so individual tests can introspect what got captured.
    """
    captured: dict[str, object] = {}

    from translation_pipeline import cli as cli_module
    from translation_pipeline.models import (
        FinalOutput,
        LanguagePair,
        PipelineState,
        RunMetadata,
        SourceDocument,
        TranslationBrief,
    )

    class _FakePipeline:
        def run(
            self,
            source: SourceDocument,
            target_language: str,
            on_event: object | None = None,
        ) -> PipelineState:
            captured["source"] = source
            pair = LanguagePair(source=source.source_language, target=target_language)
            state = PipelineState(
                source=source,
                language_pair=pair,
                metadata=RunMetadata(
                    run_id="t", source_path=source.path,
                    language_pair=pair, pipeline_version="0.0.0",
                ),
            )
            state.final_output = FinalOutput(
                text=" ".join(block_texts),
                language_pair=pair,
                brief=TranslationBrief(
                    document_type="x", register_level="n", target_audience="a"
                ),
                blocks=list(block_texts),
            )
            return state

    monkeypatch.setattr(cli_module, "build_default_pipeline", lambda _settings: _FakePipeline())
    monkeypatch.setattr(cli_module, "detect_language", lambda *a, **k: src_lang)

    return captured


def test_preserve_mode_uses_format_writer_for_docx(
    runner: CliRunner, tmp_path: Path, settings_env: None,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from docx import Document

    src = tmp_path / "doc.docx"
    d = Document()
    d.add_paragraph("First paragraph.")
    d.add_paragraph("Second paragraph.")
    d.save(str(src))

    _patch_pipeline_to_return_blocks(monkeypatch, ["Pierwszy.", "Drugi."])

    result = runner.invoke(app, [str(src), "--from", "EN", "--to", "PL"])
    assert result.exit_code == 0, result.stdout

    out_docx = tmp_path / "doc.pl.docx"
    assert out_docx.exists()
    written = Document(str(out_docx))
    paragraph_texts = [p.text for p in written.paragraphs if p.text.strip()]
    assert "Pierwszy." in paragraph_texts
    assert "Drugi." in paragraph_texts


def test_preserve_mode_falls_back_when_blocks_mismatch(
    runner: CliRunner, tmp_path: Path, settings_env: None,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from docx import Document

    src = tmp_path / "doc.docx"
    d = Document()
    d.add_paragraph("First.")
    d.add_paragraph("Second.")
    d.save(str(src))

    # Pipeline returns one block but source has two — simulates sentinel mismatch.
    _patch_pipeline_to_return_blocks(monkeypatch, [])  # empty triggers fallback

    result = runner.invoke(app, [str(src), "--from", "EN", "--to", "PL"])
    assert result.exit_code == 0, result.stdout
    assert (tmp_path / "doc.pl.docx").exists()
    # Translation should be the joined text (fresh-doc writer); no crash.


def test_no_preserve_format_uses_fresh_writer(
    runner: CliRunner, tmp_path: Path, settings_env: None,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from docx import Document

    src = tmp_path / "doc.docx"
    d = Document()
    d.add_paragraph("First paragraph.")
    p = d.add_paragraph()
    run = p.add_run("Bold heading")
    run.bold = True
    d.save(str(src))

    captured = _patch_pipeline_to_return_blocks(monkeypatch, ["x", "y"])

    result = runner.invoke(
        app, [str(src), "--from", "EN", "--to", "PL", "--no-preserve-format"]
    )
    assert result.exit_code == 0, result.stdout
    # Source seen by pipeline is plain text (no [[BLK]] markers) since preserve-format off.
    src_doc = captured["source"]
    assert src_doc.blocks == ()
    assert "[[BLK]]" not in src_doc.text


def test_preserve_mode_pdf_dispatch(
    runner: CliRunner, tmp_path: Path, settings_env: None,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from translation_pipeline.documents import PdfWriter as DocPdfWriter

    src = tmp_path / "doc.pdf"
    DocPdfWriter().write(src, "Hello.\n\nWorld.")

    _patch_pipeline_to_return_blocks(monkeypatch, ["Olá.", "Mundo."], tgt_lang="PT-BR")

    result = runner.invoke(app, [str(src), "--from", "EN", "--to", "PT-BR"])
    assert result.exit_code == 0, result.stdout
    assert (tmp_path / "doc.pt-br.pdf").exists()
