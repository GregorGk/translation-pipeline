"""Document I/O — text extraction + format-preserving output for .docx / .pdf.

Two reading modes:

- ``read_document(path)`` — flat text, used by the txt path and ``--no-preserve-format``.
- ``read_structured(path)`` — returns a ``StructuredSource`` with an ordered list of
  ``SourceBlock`` records. Each block carries the source text and an opaque ``anchor``
  pointing back into the live source document (a python-docx run-list, or a PyMuPDF
  page rect). The CLI joins block text via ``[[BLK]]`` sentinels for the pipeline,
  splits the translated output back on the same sentinel, and reinjects each
  translated block into its anchor via ``DocxFormatPreservingWriter`` /
  ``PdfFormatPreservingWriter``.

Two writing modes:

- ``DocxFormatPreservingWriter`` / ``PdfFormatPreservingWriter`` — mutate the original
  document in place. Output is visually indistinguishable from source except for the
  translated text. Used by default for .docx and .pdf inputs.
- ``DocxWriter`` / ``PdfWriter`` — fresh-document fallback. Used for .txt input, for
  ``--no-preserve-format``, or when sentinel alignment fails.

PDF format preservation uses PyMuPDF (AGPL-3.0). Outputs from this tool inherit AGPL
unless a commercial PyMuPDF license is purchased — surfaced once at first use.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Literal

log = logging.getLogger(__name__)

SUPPORTED_INPUT_EXTENSIONS: tuple[str, ...] = (".txt", ".pdf", ".docx")

SENTINEL = "[[BLK]]"
JOINED_SENTINEL = f"\n\n{SENTINEL}\n\n"
# When source text already contains the literal sentinel, escape with a zero-width
# space so it round-trips through the pipeline without confusing the splitter.
_ESCAPED_SENTINEL = "[[BLK​]]"


class UnsupportedDocumentFormat(Exception):
    """Raised when the input file extension isn't one we can read."""


# ---- Block model -------------------------------------------------------------


@dataclass(frozen=True)
class SourceBlock:
    index: int
    text: str
    # Opaque anchor:
    #   .docx — list of <w:r> XML elements making up the paragraph or cell
    #   .pdf  — (page_index: int, rect: fitz.Rect, font_size: float)
    #   .txt  — None
    anchor: Any = None


@dataclass(frozen=True)
class StructuredSource:
    kind: Literal["txt", "docx", "pdf"]
    blocks: tuple[SourceBlock, ...]
    # Live document handle: docx.Document, fitz.Document, or None for .txt.
    handle: Any = None
    # Raw character count of the source (for cost estimation parity with old reader).
    raw_chars: int = 0
    extras: dict[str, Any] = field(default_factory=dict)

    @property
    def joined_text(self) -> str:
        """Block texts joined with ``[[BLK]]`` sentinels."""
        if not self.blocks:
            return ""
        return JOINED_SENTINEL.join(_escape_sentinel(b.text) for b in self.blocks)


def _escape_sentinel(text: str) -> str:
    if SENTINEL in text:
        return text.replace(SENTINEL, _ESCAPED_SENTINEL)
    return text


def _unescape_sentinel(text: str) -> str:
    if _ESCAPED_SENTINEL in text:
        return text.replace(_ESCAPED_SENTINEL, SENTINEL)
    return text


# ---- Public read API ---------------------------------------------------------


def read_document(path: Path) -> str:
    """Extract plain text from a supported document type (no anchors).

    Kept for ``--no-preserve-format`` and ``--dry-run`` paths. Format-preserving
    runs use ``read_structured`` instead.
    """
    ext = path.suffix.lower()
    if ext == ".txt":
        return _read_txt(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    raise UnsupportedDocumentFormat(
        f"unsupported input extension {ext!r}; supported: {', '.join(SUPPORTED_INPUT_EXTENSIONS)}"
    )


def read_structured(path: Path) -> StructuredSource:
    """Extract a block-anchored view of a supported document type."""
    ext = path.suffix.lower()
    if ext == ".txt":
        return _read_txt_structured(path)
    if ext == ".pdf":
        return _read_pdf_structured(path)
    if ext == ".docx":
        return _read_docx_structured(path)
    raise UnsupportedDocumentFormat(
        f"unsupported input extension {ext!r}; supported: {', '.join(SUPPORTED_INPUT_EXTENSIONS)}"
    )


# ---- TXT ---------------------------------------------------------------------


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _read_txt_structured(path: Path) -> StructuredSource:
    text = _read_txt(path)
    blocks = tuple(
        SourceBlock(index=i, text=p.strip(), anchor=None)
        for i, p in enumerate(text.split("\n\n"))
        if p.strip()
    )
    return StructuredSource(kind="txt", blocks=blocks, handle=None, raw_chars=len(text))


# ---- PDF (read) --------------------------------------------------------------


def _read_pdf(path: Path) -> str:
    """Plain-text PDF extraction via pypdf — used by --no-preserve-format / --dry-run."""
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(path))
    except PdfReadError as e:
        raise OSError(f"could not read PDF {path}: {e}") from e

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as e:
            raise OSError(
                f"PDF {path} is encrypted; decryption with empty password failed: {e}"
            ) from e

    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as e:
            raise OSError(f"failed to extract text from page {i + 1} of {path}: {e}") from e

    return "\n\n".join(p.strip() for p in pages if p.strip())


def _read_pdf_structured(path: Path, *, max_y_gap_factor: float = 2.5) -> StructuredSource:
    """Structured PDF extraction via PyMuPDF — yields per-paragraph anchors.

    Granular extraction is at the **span** level (PyMuPDF dict mode). A span is
    PDF's analog of a Word run: a contiguous text fragment with identical font,
    size, and color. We then coarsen vertically-adjacent spans on the same
    page into logical paragraphs so the LLM only has to preserve a sane number
    of ``[[BLK]]`` sentinels. Each block's anchor is a list of
    ``(rect, font_size, span_text)`` triples plus an aggregate rect.
    """
    import fitz  # PyMuPDF

    try:
        doc = fitz.open(str(path))
    except Exception as e:
        raise OSError(f"could not read PDF {path}: {e}") from e

    if doc.is_encrypted and not doc.authenticate(""):
        raise OSError(
            f"PDF {path} is encrypted; decryption with empty password failed"
        )

    raw_units: list[tuple[int, Any, str, float]] = []
    raw_chars = 0
    for page_idx in range(doc.page_count):
        page = doc.load_page(page_idx)
        for rect, text, font_size in _pdf_spans(page):
            raw_units.append((page_idx, rect, text, font_size))
            raw_chars += len(text)

    blocks = _coarsen_pdf_units(raw_units, max_y_gap_factor=max_y_gap_factor)
    return StructuredSource(
        kind="pdf", blocks=tuple(blocks), handle=doc, raw_chars=raw_chars
    )


def _pdf_spans(page: Any) -> list[tuple[Any, str, float]]:
    """Yield ``(rect, text, font_size)`` per span via dict-mode extraction.

    Adjacent spans within the same line that share font_size are merged so
    we don't fragment a sentence into individual word-spans (PyMuPDF
    sometimes emits one span per word with ligatures).
    """
    import fitz

    out: list[tuple[Any, str, float]] = []
    info = page.get_text("dict")
    for block in info.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            current: tuple[Any, str, float] | None = None
            for span in line.get("spans", []):
                text = span.get("text") or ""
                if not text.strip():
                    if current is not None:
                        # Whitespace span between same-styled neighbours: append
                        # to current to preserve original spacing.
                        rect, t, sz = current
                        bbox = span.get("bbox") or [rect.x0, rect.y0, rect.x1, rect.y1]
                        new_rect = fitz.Rect(rect.x0, rect.y0, bbox[2], rect.y1)
                        current = (new_rect, t + text, sz)
                    continue
                bbox = span.get("bbox")
                if not bbox:
                    continue
                size = float(span.get("size") or 11.0)
                rect = fitz.Rect(*bbox)
                if (
                    current is not None
                    and abs(current[2] - size) < 0.1
                    and abs(rect.y0 - current[0].y0) < 1.5
                ):
                    # Same line, same size — merge.
                    merged_rect = fitz.Rect(current[0])
                    merged_rect.include_rect(rect)
                    current = (merged_rect, current[1] + text, size)
                else:
                    if current is not None:
                        out.append(current)
                    current = (rect, text, size)
            if current is not None:
                out.append(current)
    return out


def _coarsen_pdf_units(
    units: list[tuple[int, Any, str, float]],
    *,
    max_y_gap_factor: float,
) -> list[SourceBlock]:
    """Group consecutive PDF text units (spans) into paragraph-level blocks.

    Two units belong to the same group when:
      - they're on the same page, AND
      - the vertical gap between them is small (< ``max_y_gap_factor`` ×
        the average font size of the previous unit).

    Each output block carries text joined by space (so words flow naturally
    across line wraps), plus an anchor
    ``(page_idx, [(rect, text, font_size), ...], aggregate_rect, font_size)``.
    The list of sub-units carries per-span text so the writer can redistribute
    the translated paragraph back across spans (preserving inline font / size
    changes).
    """
    import fitz

    groups: list[list[tuple[int, Any, str, float]]] = []
    for unit in units:
        page_idx, rect, _text, font_size = unit
        if groups:
            last = groups[-1][-1]
            last_page, last_rect, _last_text, last_size = last
            line_height = max(last_size, 8.0)
            if (
                page_idx == last_page
                and rect.y0 - last_rect.y1 <= line_height * max_y_gap_factor
            ):
                groups[-1].append(unit)
                continue
        groups.append([unit])

    out: list[SourceBlock] = []
    for group in groups:
        page_idx = group[0][0]
        sub_units = [(u[1], u[2], u[3]) for u in group]
        text = " ".join(u[2].strip() for u in group if u[2].strip())
        sizes = [u[3] for u in group if u[3] > 0]
        font_size = sum(sizes) / len(sizes) if sizes else 11.0
        agg = fitz.Rect(sub_units[0][0])
        for r, _t, _s in sub_units[1:]:
            agg.include_rect(r)
        out.append(
            SourceBlock(
                index=len(out),
                text=text,
                anchor=(page_idx, sub_units, agg, font_size),
            )
        )
    return out


def _estimate_font_size(page: Any, rect: Any) -> float:
    """Best-effort font-size hint pulled from the dict-form text spans inside ``rect``."""
    try:
        info = page.get_text("dict", clip=rect)
    except Exception:
        return 11.0
    sizes: list[float] = []
    for block in info.get("blocks", []):
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                size = span.get("size")
                if isinstance(size, (int, float)):
                    sizes.append(float(size))
    if not sizes:
        return 11.0
    return sum(sizes) / len(sizes)


# ---- DOCX (read) -------------------------------------------------------------


def _read_docx(path: Path) -> str:
    """Plain-text DOCX extraction — used by --no-preserve-format / --dry-run."""
    from docx import Document
    from docx.oxml.ns import qn

    try:
        doc = Document(str(path))
    except Exception as e:
        raise OSError(f"could not read DOCX {path}: {e}") from e

    parts: list[str] = []
    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if text.strip():
                parts.append(text)
        elif tag == qn("w:tbl"):
            for row in child.iter(qn("w:tr")):
                row_texts: list[str] = []
                for cell in row.iter(qn("w:tc")):
                    cell_text = "".join(t.text or "" for t in cell.iter(qn("w:t")))
                    if cell_text.strip():
                        row_texts.append(cell_text.strip())
                if row_texts:
                    parts.append(" | ".join(row_texts))

    return "\n\n".join(parts)


def _read_docx_structured(path: Path) -> StructuredSource:
    """Structured DOCX extraction — paragraphs + table cells + headers/footers/footnotes.

    Each block's anchor is the list of ``<w:r>`` run elements that hold its text,
    so the writer can rewrite them in place.
    """
    from docx import Document
    from docx.oxml.ns import qn

    try:
        doc = Document(str(path))
    except Exception as e:
        raise OSError(f"could not read DOCX {path}: {e}") from e

    blocks: list[SourceBlock] = []
    raw_chars = 0

    def emit_paragraph(p_element: Any) -> None:
        nonlocal raw_chars
        runs = list(p_element.iter(qn("w:r")))
        if not runs:
            return
        text = "".join(_run_text(r) for r in runs)
        if not text.strip():
            return
        # Filter out runs that hold drawings only (no <w:t>) — keep only text-bearing
        # runs as anchors so the writer doesn't blow away non-text runs.
        text_runs = [r for r in runs if list(r.iter(qn("w:t")))]
        if not text_runs:
            return
        blocks.append(
            SourceBlock(index=len(blocks), text=text, anchor=text_runs)
        )
        raw_chars += len(text)

    def emit_paragraphs_in(container: Any) -> None:
        for child in container.iterchildren():
            tag = child.tag
            if tag == qn("w:p"):
                emit_paragraph(child)
            elif tag == qn("w:tbl"):
                for cell in child.iter(qn("w:tc")):
                    for cell_child in cell.iterchildren():
                        if cell_child.tag == qn("w:p"):
                            emit_paragraph(cell_child)

    # 1. Body content
    emit_paragraphs_in(doc.element.body)

    # 2. Headers and footers from every section. Some properties may raise on
    # documents without them — suppress and move on.
    import contextlib

    for section in doc.sections:
        for hf in (
            section.header, section.footer, section.first_page_header,
            section.first_page_footer, section.even_page_header,
            section.even_page_footer,
        ):
            with contextlib.suppress(Exception):
                emit_paragraphs_in(hf._element)

    # 3. Footnotes part if present. python-docx exposes the related part but its
    # element accessor varies across versions; treat any failure as "no footnotes."
    with contextlib.suppress(Exception):
        footnotes_part = doc.part.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
        )
        root = getattr(footnotes_part, "element", None) or getattr(footnotes_part, "_element", None)
        if root is not None:
            for footnote in root.iter(qn("w:footnote")):
                for p in footnote.iter(qn("w:p")):
                    emit_paragraph(p)

    return StructuredSource(
        kind="docx", blocks=tuple(blocks), handle=doc, raw_chars=raw_chars
    )


def _run_text(run_element: Any) -> str:
    from docx.oxml.ns import qn

    return "".join(t.text or "" for t in run_element.iter(qn("w:t")))


# ---- DOCX (fresh-document writer, fallback) ---------------------------------


class DocxWriter:
    """Fallback: write the translation as a fresh paragraph-per-block .docx file."""

    def write(self, output_path: Path, text: str) -> Path:
        from docx import Document

        doc = Document()
        for block in text.split("\n\n"):
            block = block.strip()
            if block:
                doc.add_paragraph(_unescape_sentinel(block))
        doc.save(str(output_path))
        return output_path


# ---- Run / span redistribution -----------------------------------------------


def _redistribute_to_units(
    block_translated_text: str,
    source_unit_texts: list[str],
    *,
    client: Any | None = None,
    settings: Any | None = None,
    warnings: list[str] | None = None,
    block_index: int = -1,
) -> list[str]:
    """Split a paragraph-level translation across the original sub-units (runs / spans).

    - 0 source units → empty list.
    - 1 source unit → return the full translated text in a 1-element list.
    - 2+ units → call Claude to align if a client is provided; fall back to a
      character-proportional split on alignment failure.

    Whitespace-only / empty source units always pass through untouched (we don't
    waste an alignment slot on them — they keep their literal whitespace).
    """
    n = len(source_unit_texts)
    if n == 0:
        return []
    if n == 1:
        return [block_translated_text]

    # Identify text-carrying units; whitespace-only ones get passed through.
    indexed = [(i, t) for i, t in enumerate(source_unit_texts)]
    text_carriers = [(i, t) for i, t in indexed if t.strip()]
    if len(text_carriers) <= 1:
        # Only one unit actually carries text — give the translation to it,
        # leave the rest as their original whitespace / punctuation.
        out = list(source_unit_texts)
        if text_carriers:
            target_idx = text_carriers[0][0]
            out[target_idx] = block_translated_text
        else:
            out[0] = block_translated_text
        return out

    # Multi-carrier: try LLM alignment first.
    if client is not None and settings is not None:
        try:
            aligned = _align_via_claude(
                block_translated_text,
                [t for _, t in text_carriers],
                client=client,
                settings=settings,
            )
        except Exception as e:
            log.warning("alignment call failed for block %d: %s", block_index, e)
            aligned = None
        if aligned is not None and len(aligned) == len(text_carriers):
            out = list(source_unit_texts)
            for (src_idx, _), translated in zip(text_carriers, aligned, strict=True):
                out[src_idx] = translated
            return out
        if warnings is not None:
            warnings.append(
                f"alignment count mismatch on block {block_index} — "
                "proportional split applied"
            )

    # Fallback: character-proportional split across text-carrying units.
    return _proportional_split(block_translated_text, source_unit_texts)


def _proportional_split(
    translated: str, source_unit_texts: list[str]
) -> list[str]:
    """Distribute ``translated`` across source units weighted by source length.

    Pre-computes cut positions from cumulative source-length proportions so
    every text-carrying unit ends up with at least one character. Word
    boundaries within ±6 chars of each computed cut are preferred so we don't
    slice mid-word. Whitespace-only units are echoed back verbatim.
    """
    out: list[str] = []
    text_indices = [i for i, t in enumerate(source_unit_texts) if t.strip()]
    if not text_indices:
        return list(source_unit_texts)

    text_lens = [len(source_unit_texts[i]) for i in text_indices]
    total = sum(text_lens)
    if total == 0:
        return list(source_unit_texts)

    n = len(text_indices)
    n_chars = len(translated)

    # Pre-compute cut positions for the n text-carrying units.
    cum_lens = [sum(text_lens[: k + 1]) for k in range(n)]
    raw_cuts = [round(c / total * n_chars) for c in cum_lens]
    raw_cuts[-1] = n_chars  # last carrier always closes out

    # Snap cuts forward to next whitespace (within 6 chars), keeping each cut
    # strictly greater than the previous so no carrier collapses to empty.
    snapped: list[int] = []
    for k in range(n - 1):
        prev_floor = (snapped[-1] + 1) if snapped else 1
        target = max(raw_cuts[k], prev_floor)
        # Forward scan only (don't pull a cut back into the previous slice).
        for j in range(target, min(target + 6, n_chars)):
            if translated[j].isspace():
                target = j
                break
        snapped.append(min(target, n_chars - (n - k - 1)))  # leave room
    snapped.append(n_chars)

    # Build per-carrier slices.
    starts = [0] + [snapped[k] for k in range(n - 1)]
    ends = snapped
    carrier_slices = {
        text_indices[k]: translated[starts[k] : ends[k]].strip()
        for k in range(n)
    }

    for i, src in enumerate(source_unit_texts):
        if not src.strip():
            out.append(src)
        else:
            out.append(carrier_slices[i])
    return out


def _align_via_claude(
    translated_block: str,
    source_fragments: list[str],
    *,
    client: Any,
    settings: Any,
) -> list[str] | None:
    """Single Anthropic tool_use call to align translated text across fragments."""
    from pydantic import BaseModel, ConfigDict

    from translation_pipeline.llm import anthropic_structured

    class _Aligned(BaseModel):
        model_config = ConfigDict(extra="forbid")

        translations: list[str]

    schema: dict[str, Any] = {
        "type": "object",
        "properties": {
            "translations": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["translations"],
    }

    numbered = "\n".join(f"  {i + 1}. {repr(s)}" for i, s in enumerate(source_fragments))
    prompt = _ALIGNMENT_PROMPT.render(
        numbered_source_fragments=numbered,
        translated_block=translated_block,
    )
    try:
        result = anthropic_structured(
            client,
            model=settings.MODEL_DIVERGENCE_DETECTION,
            max_tokens=4096,
            prompt=prompt,
            tool_name="submit_alignment",
            tool_description="Submit the translated fragments aligned to source.",
            tool_schema=schema,
            schema_model=_Aligned,
        )
    except Exception:
        return None
    return result.parsed.translations


# Lazy-loaded prompt (avoid importing prompts module at top-level).
class _LazyAlignmentPrompt:
    _cached: Any = None

    def render(self, **kwargs: Any) -> str:
        if self._cached is None:
            from translation_pipeline.prompts import load_prompt

            self._cached = load_prompt("run_alignment")
        rendered: str = self._cached.render(**kwargs)
        return rendered


_ALIGNMENT_PROMPT = _LazyAlignmentPrompt()


# ---- DOCX (format-preserving writer) ----------------------------------------


class DocxFormatPreservingWriter:
    """Mutate the source ``Document`` in place at run granularity.

    For each block, the source paragraph's runs each keep their own ``<w:rPr>``
    (bold, italic, font, color, hyperlink wrappers, etc.). The translated text
    gets aligned back to the original runs via a small Claude call when the
    paragraph has 2+ text-carrying runs; single-run paragraphs skip the call.
    Tables, headers/footers, lists, numbering, page geometry, images, and
    section properties are preserved by python-docx itself.
    """

    def __init__(
        self,
        client: Any | None = None,
        settings: Any | None = None,
    ) -> None:
        self._client = client
        self._settings = settings

    def write(
        self,
        output_path: Path,
        source: StructuredSource,
        translated_blocks: list[str],
        *,
        warnings: list[str] | None = None,
    ) -> Path:
        if source.kind != "docx":
            raise ValueError(f"DocxFormatPreservingWriter requires kind='docx', got {source.kind!r}")
        if len(translated_blocks) != len(source.blocks):
            raise ValueError(
                f"block count mismatch: source={len(source.blocks)}, "
                f"translated={len(translated_blocks)}"
            )

        for idx, (src_block, translated) in enumerate(
            zip(source.blocks, translated_blocks, strict=True)
        ):
            runs: list[Any] = list(src_block.anchor or [])
            if not runs:
                continue
            unescaped = _unescape_sentinel(translated.strip())
            source_unit_texts = [_run_text(r) for r in runs]
            redistributed = _redistribute_to_units(
                unescaped,
                source_unit_texts,
                client=self._client,
                settings=self._settings,
                warnings=warnings,
                block_index=idx,
            )
            for run, new_text in zip(runs, redistributed, strict=True):
                _set_run_text(run, new_text)

        source.handle.save(str(output_path))
        return output_path


def _set_run_text(run_element: Any, text: str) -> None:
    """Replace a run's text by clearing existing <w:t> elements and adding one new one."""
    from docx.oxml.ns import qn
    from lxml import etree

    # Remove existing <w:t> children (preserves <w:rPr>).
    for t in list(run_element.iter(qn("w:t"))):
        t.getparent().remove(t)
    if text == "":
        return
    new_t = etree.SubElement(run_element, qn("w:t"))
    new_t.text = text
    # Preserve leading/trailing whitespace
    new_t.set(qn("xml:space"), "preserve")


# ---- PDF (fresh-document writer, fallback via reportlab) --------------------


_UNICODE_FONT_CANDIDATES: tuple[Path, ...] = (
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    Path("/Library/Fonts/Arial Unicode.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
    Path("/usr/share/fonts/TTF/DejaVuSans.ttf"),
    Path("/usr/share/fonts/noto/NotoSans-Regular.ttf"),
    Path("/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
)


def _find_unicode_font() -> Path | None:
    for p in _UNICODE_FONT_CANDIDATES:
        if p.exists():
            return p
    return None


class PdfWriter:
    """Fallback: paragraph-per-block PDF via reportlab."""

    _FALLBACK_FONT_NAME = "Helvetica"
    _UNICODE_FONT_NAME = "TranslationPipelineUnicode"

    def write(self, output_path: Path, text: str) -> Path:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFError, TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        font_name = self._FALLBACK_FONT_NAME
        unicode_font = _find_unicode_font()
        if unicode_font is not None:
            try:
                pdfmetrics.registerFont(TTFont(self._UNICODE_FONT_NAME, str(unicode_font)))
                font_name = self._UNICODE_FONT_NAME
            except TTFError as e:
                log.warning(
                    "could not register Unicode font %s: %s — falling back to Helvetica",
                    unicode_font,
                    e,
                )
        else:
            log.warning(
                "no Unicode TTF found on system; PDF output uses Helvetica/WinAnsi "
                "(Latin only). Cyrillic / Greek glyphs may render as boxes."
            )

        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "TranslationBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=14,
            spaceAfter=8,
        )

        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=LETTER,
            leftMargin=54,
            rightMargin=54,
            topMargin=54,
            bottomMargin=54,
            title=output_path.stem,
        )
        story: list[object] = []
        for block in text.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            block = _unescape_sentinel(block)
            safe = (
                block.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace("\n", "<br/>")
            )
            story.append(Paragraph(safe, body))
            story.append(Spacer(1, 4))
        if not story:
            story.append(Paragraph("(empty)", body))
        doc.build(story)
        return output_path


# ---- PDF (format-preserving writer via PyMuPDF) -----------------------------


_AGPL_NOTICE_LOGGED = False


def _unpack_pdf_anchor(
    anchor: tuple[Any, ...],
) -> tuple[list[Any], list[float], list[str], int, Any, float]:
    """Normalise the various PDF anchor shapes the readers have produced.

    Returns ``(sub_rects, sub_sizes, source_unit_texts, page_idx, agg, paragraph_size)``.
    Supports:
      - legacy 3-tuple ``(page_idx, rect, font_size)``
      - rect-list 4-tuple ``(page_idx, [rects], agg, font_size)``
      - span-aware 4-tuple ``(page_idx, [(rect, text, size), ...], agg, font_size)``
    """
    if len(anchor) == 3:
        page_idx, rect, font_size = anchor
        return [rect], [font_size], [""], page_idx, rect, font_size
    page_idx, sub, agg, font_size = anchor
    if not sub:
        return [], [], [], page_idx, agg, font_size
    first = sub[0]
    if isinstance(first, tuple) and len(first) == 3:
        rects = [u[0] for u in sub]
        texts = [u[1] for u in sub]
        sizes = [u[2] for u in sub]
        return rects, sizes, texts, page_idx, agg, font_size
    # Plain list of rects (legacy)
    return list(sub), [font_size] * len(sub), [""] * len(sub), page_idx, agg, font_size


class PdfFormatPreservingWriter:
    """In-place PDF text replacement via PyMuPDF, span-aware.

    For each source paragraph: redact each original span rect, then for every
    span insert the corresponding aligned-translation fragment back into the
    same rect with the same font size. Multi-span paragraphs invoke a small
    Claude alignment call so each span keeps its own font / size in the
    target language. Single-span paragraphs skip the call.

    On overflow, font size is shrunk; if still overflowing, the rect expands
    downward by 20% once. Images, vector graphics, page geometry are preserved.
    """

    _FONT_NAME_BUILTIN = "helv"
    _FONT_NAME_UNICODE = "tp_unicode"
    _MIN_FONT_SIZE = 6.0
    _SHRINK_STEP = 0.5

    def __init__(
        self,
        client: Any | None = None,
        settings: Any | None = None,
    ) -> None:
        self._client = client
        self._settings = settings

    def write(
        self,
        output_path: Path,
        source: StructuredSource,
        translated_blocks: list[str],
        *,
        warnings: list[str] | None = None,
    ) -> Path:
        global _AGPL_NOTICE_LOGGED
        if not _AGPL_NOTICE_LOGGED:
            log.warning(
                "PyMuPDF is AGPL-3.0; outputs from this tool inherit AGPL unless you "
                "obtain a commercial PyMuPDF license."
            )
            _AGPL_NOTICE_LOGGED = True

        if source.kind != "pdf":
            raise ValueError(f"PdfFormatPreservingWriter requires kind='pdf', got {source.kind!r}")
        if len(translated_blocks) != len(source.blocks):
            raise ValueError(
                f"block count mismatch: source={len(source.blocks)}, "
                f"translated={len(translated_blocks)}"
            )

        doc = source.handle
        unicode_font_path = _find_unicode_font()
        # Pre-register the Unicode font once per page so insert_textbox can
        # reference it by name; fontfile= on insert_textbox would re-embed on
        # every call, bloating the output.
        if unicode_font_path is not None:
            for page_idx in range(doc.page_count):
                try:
                    doc.load_page(page_idx).insert_font(
                        fontname=self._FONT_NAME_UNICODE,
                        fontfile=str(unicode_font_path),
                    )
                except Exception as e:
                    log.warning(
                        "could not embed Unicode font on page %d: %s — "
                        "PDF will use built-in helv (Latin only)",
                        page_idx,
                        e,
                    )
                    unicode_font_path = None
                    break

        font_name = (
            self._FONT_NAME_UNICODE if unicode_font_path is not None
            else self._FONT_NAME_BUILTIN
        )
        font_file = str(unicode_font_path) if unicode_font_path is not None else None

        # Anchor schema: (page_idx, [(rect, span_text, font_size), ...],
        # aggregate_rect, paragraph_font_size). Legacy 3-tuple anchors
        # (page_idx, rect, font_size) and rect-only-list 4-tuples are still
        # accepted for backward compat with older test fixtures.
        # Per-block redistribution: each multi-span block routes through Claude
        # to align translated text back to per-span fragments.
        by_page: dict[int, list[tuple[list[Any], list[float], list[str], Any, float]]] = {}
        for idx, (src_block, translated) in enumerate(
            zip(source.blocks, translated_blocks, strict=True)
        ):
            anchor = src_block.anchor
            if anchor is None:
                continue
            sub_rects, sub_sizes, source_unit_texts, page_idx, agg, para_size = (
                _unpack_pdf_anchor(anchor)
            )
            unescaped = _unescape_sentinel(translated.strip())
            redistributed = _redistribute_to_units(
                unescaped,
                source_unit_texts,
                client=self._client,
                settings=self._settings,
                warnings=warnings,
                block_index=idx,
            )
            by_page.setdefault(page_idx, []).append(
                (sub_rects, sub_sizes, redistributed, agg, para_size)
            )

        for page_idx, items in by_page.items():
            page = doc.load_page(page_idx)
            for sub_rects, _sub_sizes, _texts, _agg, _ps in items:
                for r in sub_rects:
                    page.add_redact_annot(r, fill=(1, 1, 1))
            page.apply_redactions()
            for sub_rects, sub_sizes, texts, _agg, _ps in items:
                for rect, size, fragment in zip(
                    sub_rects, sub_sizes, texts, strict=True
                ):
                    if not fragment.strip():
                        continue
                    self._insert_with_overflow(
                        page, rect, fragment, size, font_name, font_file
                    )

        doc.save(str(output_path), garbage=4, deflate=True)
        return output_path

    def _insert_with_overflow(
        self,
        page: Any,
        rect: Any,
        text: str,
        font_size: float,
        font_name: str,
        font_file: str | None,
    ) -> None:
        import fitz

        size = max(font_size, self._MIN_FONT_SIZE)
        kwargs: dict[str, Any] = {"fontname": font_name, "align": 0}
        if font_file is not None:
            # Pass on every call so PyMuPDF can resolve missing-font references on
            # pages where insert_font wasn't applied (defensive — should be rare).
            kwargs["fontfile"] = font_file

        while size >= self._MIN_FONT_SIZE:
            try:
                result = page.insert_textbox(rect, text, fontsize=size, **kwargs)
            except Exception:
                # Font lookup glitch: drop fontfile and retry once with builtin.
                kwargs.pop("fontfile", None)
                kwargs["fontname"] = self._FONT_NAME_BUILTIN
                result = page.insert_textbox(rect, text, fontsize=size, **kwargs)
            if result >= 0:
                return
            size -= self._SHRINK_STEP

        # Last resort: expand rect downward by 20% and retry at the floor size.
        expanded = fitz.Rect(
            rect.x0, rect.y0, rect.x1, rect.y1 + (rect.height * 0.2)
        )
        try:
            page.insert_textbox(
                expanded, text, fontsize=self._MIN_FONT_SIZE, **kwargs
            )
        except Exception:
            kwargs.pop("fontfile", None)
            kwargs["fontname"] = self._FONT_NAME_BUILTIN
            page.insert_textbox(
                expanded, text, fontsize=self._MIN_FONT_SIZE, **kwargs
            )
